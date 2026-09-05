# scripts/build_draft_template.py
"""Generate the .docx templates the drafting worker fills.

Run: .venv/bin/python -m scripts.build_draft_template

The templates are checked in as .docx, because that is the point of using
docxtpl: the layout lives in a file a person opens in Word and edits, with
no code deployment. This script exists so the first version is reproducible
and reviewable as a diff, not so it stays the only way to change them. Once
someone opens `s138_demand_notice.docx` in Word and moves a margin, this
script is history and the .docx is the source.

No letterhead, deliberately. A s.138 notice goes out on the advocate's own
letterhead carrying their Bar Council enrolment number, signature and seal;
we hold none of that and must not imply it. What we produce is a draft an
advocate settles, which is why the footer says so and why the drafting
notes sit at the end for them to delete.
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

TEMPLATES = Path(__file__).resolve().parents[1] / "src" / "legal_ai" / "drafting" / "templates"

# Courts and opposing counsel read these on paper. 12pt serif with generous
# leading is the register; a UI font would read as a printout of a web page.
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)


def _styled(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.4

    for section in document.sections:
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)


def _para(document: Document, text: str = "", *, bold=False, align=None, size=None):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    if size is not None:
        run.font.size = size
    if align is not None:
        paragraph.alignment = align
    return paragraph


def s138_demand_notice() -> Document:
    document = Document()
    _styled(document)

    _para(document, "NOTICE UNDER SECTION 138 OF THE NEGOTIABLE INSTRUMENTS ACT, 1881",
          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(document, "Date: {{ date }}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    _para(document, "To,", bold=True)
    _para(document, "{{ recipient.name }}")
    _para(document, "{{ recipient.address }}")

    _para(document, "Subject: {{ subject }}", bold=True)
    _para(document, "Sir,")
    _para(document,
          "Under instructions from and on behalf of my client, "
          "{{ sender.name }}, I serve upon you the following notice:")

    # Paragraph-level loops: {%p %} consumes the paragraph holding the tag,
    # so the numbered items are not separated by empty lines.
    _para(document, "{%p for fact in facts %}")
    _para(document, "{{ loop.index }}. {{ fact }}")
    _para(document, "{%p endfor %}")

    _para(document, "{%p for ground in grounds %}")
    _para(document, "{{ facts|length + loop.index }}. {{ ground.text }}"
          "{% if ground.authority %} ({{ ground.authority }}){% endif %}")
    _para(document, "{%p endfor %}")

    _para(document, "DEMAND", bold=True)
    _para(document,
          "You are hereby called upon to {{ demand.what }} within "
          "{{ demand.within_days }} days of {{ demand.from_when }}.")
    _para(document, "{{ consequence }}")

    _para(document, "{%p if annexures %}")
    _para(document, "ANNEXURES", bold=True)
    _para(document, "{%p for annexure in annexures %}")
    _para(document, "{{ loop.index }}. {{ annexure }}")
    _para(document, "{%p endfor %}")
    _para(document, "{%p endif %}")

    _para(document, "Yours faithfully,")
    _para(document)
    _para(document, "_______________________________")
    _para(document, "Advocate")
    _para(document, "[Name, Bar Council enrolment number and address to be completed]")

    # Everything the draft could not settle, kept together at the end so an
    # advocate can read it once and delete the block.
    _para(document, "{%p if needs_input or warnings %}")
    document.add_page_break()
    _para(document, "DRAFTING NOTES — DELETE BEFORE SENDING", bold=True)

    _para(document, "{%p if warnings %}")
    _para(document, "Resolve before this is sent:", bold=True)
    _para(document, "{%p for warning in warnings %}")
    _para(document, "• {{ warning }}")
    _para(document, "{%p endfor %}")
    _para(document, "{%p endif %}")

    _para(document, "{%p if needs_input %}")
    _para(document, "To be supplied:", bold=True)
    _para(document, "{%p for item in needs_input %}")
    _para(document, "• {{ item }}")
    _para(document, "{%p endfor %}")
    _para(document, "{%p endif %}")
    _para(document, "{%p endif %}")

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "DRAFT prepared by Pramāṇa AI from this matter's own record. "
        "To be settled, signed and issued by an advocate on their letterhead. "
        "Verify every fact and provision before sending."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    return document


def legal_opinion() -> Document:
    """An opinion, which has no statutory form to satisfy.

    Numbered points rather than a notice's numbered paragraphs, and a
    conclusion where a notice makes a demand. Addressed to the client who
    asked, so there is no "To," block and no service details.
    """
    document = Document()
    _styled(document)

    _para(document, "LEGAL OPINION", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(document, "{{ subject }}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(document, "Date: {{ date }}", align=WD_ALIGN_PARAGRAPH.RIGHT)

    _para(document, "{%p if recipient.name %}")
    _para(document, "For: {{ recipient.name }}")
    _para(document, "{%p endif %}")

    _para(document, "{%p if facts %}")
    _para(document, "FACTS AS UNDERSTOOD", bold=True)
    _para(document, "{%p for fact in facts %}")
    _para(document, "{{ loop.index }}. {{ fact }}")
    _para(document, "{%p endfor %}")
    _para(document, "{%p endif %}")

    _para(document, "THE POSITION", bold=True)
    _para(document, "{%p for ground in grounds %}")
    _para(document, "{{ loop.index }}. {{ ground.text }}"
          "{% if ground.authority %} ({{ ground.authority }}){% endif %}")
    _para(document, "{%p endfor %}")

    _para(document, "{%p if conclusion %}")
    _para(document, "CONCLUSION", bold=True)
    _para(document, "{{ conclusion }}")
    _para(document, "{%p endif %}")

    _para(document, "{%p if needs_input or warnings %}")
    _para(document, "{%p if warnings %}")
    _para(document, "MATTERS TO RESOLVE", bold=True)
    _para(document, "{%p for warning in warnings %}")
    _para(document, "• {{ warning }}")
    _para(document, "{%p endfor %}")
    _para(document, "{%p endif %}")

    _para(document, "{%p if needs_input %}")
    _para(document, "TO ADVISE FURTHER WE WOULD NEED", bold=True)
    _para(document, "{%p for item in needs_input %}")
    _para(document, "• {{ item }}")
    _para(document, "{%p endfor %}")
    _para(document, "{%p endif %}")
    _para(document, "{%p endif %}")

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "DRAFT prepared by Pramāṇa AI from this matter's own record, on the "
        "authorities cited. Legal information, not legal advice. "
        "Verify every provision before relying on it."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run.italic = True

    return document


def main() -> None:
    TEMPLATES.mkdir(parents=True, exist_ok=True)
    for name, build in (
        ("s138_demand_notice", s138_demand_notice),
        ("legal_opinion", legal_opinion),
    ):
        path = TEMPLATES / f"{name}.docx"
        build().save(path)
        print(f"wrote {path}")


if __name__ == "__main__":
    main()
