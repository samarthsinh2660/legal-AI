"""Turn a drafted structure into a .docx.

Built directly with python-docx rather than filling a template file. There
was a template per document type, and it went wrong twice over: the
templates were binaries git never tracked, so the server pulled the feature
and had nothing to render with; and a template per instrument is a promise
to write one for every document Indian practice uses.

A document here is headings with numbered paragraphs under them, which is
what every legal document is, so one renderer draws all of them. The
headings come from the draft, not from this file.

.docx and not PDF because what we produce is a draft: the advocate puts it
on their letterhead, adds their enrolment number, settles the facts and
signs. Word is the format of the drafting stage precisely because it can be
edited; handing a lawyer a PDF hands them something to retype.
"""

from __future__ import annotations

import io
from datetime import date

from legal_ai.drafting.models import DraftStructure

# Courts and opposing counsel read these on paper. 12pt serif with generous
# leading is the register; a UI font reads as a printout of a web page.
BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 12

# What the advocate has to fill where we hold no value. Visible on purpose:
# a blank space reads as nothing missing.
UNFILLED = "__________"

FOOTER = (
    "DRAFT prepared by Pramāṇa AI from this matter's own record, on the "
    "authorities cited. To be settled and signed by an advocate. Verify "
    "every fact and provision before relying on it."
)


def render(
    draft: DraftStructure,
    *,
    today: date | None = None,
    citations: dict[str, str] | None = None,
) -> bytes:
    """The draft as a .docx, ready to download.

    `citations` maps an authority id to how it should be cited. Without one
    the body would print `act:2189:sec-138`, which is how this system
    addresses a provision and not how a document cites one -- prefer
    `render_with_citations`.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    citations = citations or {}
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.4

    section = document.sections[0]
    section.left_margin = section.right_margin = Inches(1.1)
    section.top_margin = section.bottom_margin = Inches(1.0)

    def para(text="", *, bold=False, align=None):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        if align is not None:
            paragraph.alignment = align
        return paragraph

    para(draft.title or "DRAFT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if draft.subject:
        para(draft.subject, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        f"Date: {(today or date.today()).strftime('%d %B %Y')}",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    if draft.addressed_to:
        para("To,", bold=True)
        para(draft.addressed_to)
    if draft.on_behalf_of:
        para(f"On behalf of: {draft.on_behalf_of}")

    for block in draft.sections:
        if block.heading:
            para(block.heading.upper(), bold=True)
        for number, paragraph in enumerate(block.paragraphs, start=1):
            para(f"{number}. {_with_citations(paragraph, citations)}")

    para()
    para("_______________________________")
    para("Advocate")
    para("[Name, Bar Council enrolment number and address to be completed]")

    # Everything the draft could not settle, kept together at the end so an
    # advocate reads it once and deletes the block.
    if draft.warnings or draft.needs_input:
        document.add_page_break()
        para("DRAFTING NOTES — DELETE BEFORE SENDING", bold=True)
        if draft.warnings:
            para("Resolve before this is used:", bold=True)
            for warning in draft.warnings:
                para(f"• {warning}")
        if draft.needs_input:
            para("To be supplied:", bold=True)
            for item in draft.needs_input:
                para(f"• {item}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(FOOTER)
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_with_citations(conn, draft: DraftStructure, *, today: date | None = None) -> bytes:
    """`render`, with every authority resolved against the corpus first."""
    from legal_ai.drafting.citation import format_citation

    citations = {}
    for block in draft.sections:
        for paragraph in block.paragraphs:
            for authority in paragraph.authorities:
                if authority not in citations:
                    cited = format_citation(conn, authority)
                    if cited is not None:
                        citations[authority] = cited
    return render(draft, today=today, citations=citations)


def _with_citations(paragraph, citations: dict[str, str]) -> str:
    """The paragraph, with its authorities cited after it.

    An id the corpus cannot place is dropped rather than printed: a raw
    identifier reads as a bug to whoever receives the document, and the
    proposition still stands on the text around it.
    """
    cited = [citations[a] for a in paragraph.authorities if a in citations]
    if not cited:
        return paragraph.text
    return f"{paragraph.text} ({'; '.join(cited)})"
