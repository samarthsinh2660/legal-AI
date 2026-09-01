"""Uploaded case documents: file in, text out, stored against one case.

**These do not go in `documents`.** That table is the public corpus and
`hybrid_search` reads it, so a client's pleading placed there would be
returned as authority -- and could surface for a different user's query.
Keeping private material out of the corpus table is a stronger guarantee
than remembering to filter it at every retrieval path, and there are
several.

The text is kept because the Document Agent reads it, and because a
re-extraction (a better prompt, a new field) must not require the user to
upload the file again. Nothing embeds it and nothing chunks it into
`document_chunks`.
"""

from __future__ import annotations

import io
import json
from datetime import datetime, timezone
from pathlib import Path

import psycopg

from legal_ai.ingestion.schema import content_hash

# Formats we can read text out of. A scanned PDF with no text layer parses
# to nothing; that is reported as an empty extraction rather than guessed
# at, because OCR is a different problem and pretending otherwise would
# put an empty document in a case and call it read.
SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx", ".txt", ".md"})

# Uploads are bounded so one exhibit bundle cannot exhaust memory. A
# thousand-page petition is roughly 3 MB of text, so this is generous.
MAX_TEXT_CHARS = 4_000_000


class UnsupportedDocument(Exception):
    """Raised for a file type we cannot read text from."""


def ensure_case_file_schema(conn: psycopg.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS case_files (
            document_id TEXT PRIMARY KEY,
            case_id TEXT NOT NULL REFERENCES cases(case_id) ON DELETE CASCADE,
            filename TEXT NOT NULL,
            media_type TEXT NOT NULL,
            full_text TEXT NOT NULL,
            content_hash TEXT NOT NULL,
            uploaded_at TIMESTAMPTZ NOT NULL
        )
        """
    )
    # Extracted structure, stored so opening a case does not re-run the
    # Document Agent over every file. Extraction costs a model call per
    # 12,000-character window; re-doing it on each view would make a case
    # with ten exhibits expensive to look at.
    # Catalog first: `ADD COLUMN IF NOT EXISTS` takes ACCESS EXCLUSIVE even
    # when it does nothing, and a pending request for it blocks every reader
    # behind it. See knowledge/static/db.py::ensure_bench_schema.
    has_facts = conn.execute(
        "SELECT 1 FROM information_schema.columns "
        "WHERE table_name = 'case_files' AND column_name = 'facts' AND table_schema = current_schema()"
    ).fetchone()
    if has_facts is None:
        conn.execute("ALTER TABLE case_files ADD COLUMN IF NOT EXISTS facts JSONB")
    conn.execute(
        "CREATE INDEX IF NOT EXISTS case_files_case_idx ON case_files (case_id, uploaded_at)"
    )
    conn.commit()


def extract_text(data: bytes, filename: str) -> str:
    """Text of an uploaded file, by extension.

    Returns "" when the format is readable but holds no text -- a scanned
    PDF is the common case. The caller decides what to do with that; this
    does not invent content.
    """
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise UnsupportedDocument(
            f"cannot read text from {suffix or 'a file with no extension'}; "
            f"supported: {', '.join(sorted(SUPPORTED_SUFFIXES))}"
        )

    if suffix == ".pdf":
        import pypdf

        reader = pypdf.PdfReader(io.BytesIO(data))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif suffix == ".docx":
        import docx

        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        # Legal agreements put substance in tables -- schedules, payment
        # plans, party details. Dropping them would lose exactly the facts
        # a case needs.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    else:
        text = data.decode("utf-8", errors="replace")

    return text[:MAX_TEXT_CHARS]


def store_case_file(
    conn: psycopg.Connection,
    case_id: str,
    document_id: str,
    filename: str,
    text: str,
) -> None:
    """Store an uploaded file's text against a case.

    Re-uploading the same document_id replaces its text: a user correcting
    a bad scan should not end up with two versions of one exhibit.
    """
    conn.execute(
        """
        INSERT INTO case_files (document_id, case_id, filename, media_type,
                                full_text, content_hash, uploaded_at)
        VALUES (%s, %s, %s, %s, %s, %s, %s)
        ON CONFLICT (document_id) DO UPDATE SET
            filename = EXCLUDED.filename,
            media_type = EXCLUDED.media_type,
            full_text = EXCLUDED.full_text,
            content_hash = EXCLUDED.content_hash,
            uploaded_at = EXCLUDED.uploaded_at
        """,
        (
            document_id,
            case_id,
            filename,
            Path(filename).suffix.lower().lstrip("."),
            text,
            content_hash(text),
            datetime.now(timezone.utc),
        ),
    )
    conn.commit()


def store_facts(conn: psycopg.Connection, document_id: str, facts) -> None:
    """Persist what the Document Agent extracted from a stored file.

    A failed extraction is not written: `extraction_failed` means nothing
    read the document, and storing that would make a later view treat "we
    could not read it" as "it says nothing" -- permanently, since the view
    would then find facts on record and not retry.
    """
    if getattr(facts, "extraction_failed", False):
        return
    payload = {
        "document_type": facts.document_type,
        "parties": list(facts.parties),
        "dates": list(facts.dates),
        "cited_sections": list(facts.cited_sections),
        "issues": list(facts.issues),
        "clauses": list(facts.clauses),
        "claims": list(facts.claims),
    }
    conn.execute(
        "UPDATE case_files SET facts = %s WHERE document_id = %s",
        (json.dumps(payload), document_id),
    )
    conn.commit()


def get_facts(conn: psycopg.Connection, document_id: str):
    """Stored DocumentFacts, or None if this file has not been read yet."""
    from legal_ai.context.models import DocumentFacts

    with conn.cursor() as cur:
        cur.execute("SELECT facts FROM case_files WHERE document_id = %s", (document_id,))
        row = cur.fetchone()
    if row is None or row[0] is None:
        return None
    stored = row[0]
    return DocumentFacts(
        document_id=document_id,
        document_type=stored.get("document_type"),
        parties=tuple(stored.get("parties", ())),
        dates=tuple(stored.get("dates", ())),
        cited_sections=tuple(stored.get("cited_sections", ())),
        issues=tuple(stored.get("issues", ())),
        clauses=tuple(stored.get("clauses", ())),
        claims=tuple(stored.get("claims", ())),
    )


def get_case_file_text(conn: psycopg.Connection, document_id: str) -> str | None:
    with conn.cursor() as cur:
        cur.execute("SELECT full_text FROM case_files WHERE document_id = %s", (document_id,))
        row = cur.fetchone()
    return row[0] if row else None


def list_case_files(conn: psycopg.Connection, case_id: str) -> list[tuple[str, str]]:
    """(document_id, filename) for one case, oldest first.

    Filenames only -- loading every exhibit's text to draw a file list
    would read the whole case to render a sidebar.
    """
    with conn.cursor() as cur:
        cur.execute(
            "SELECT document_id, filename FROM case_files WHERE case_id = %s "
            "ORDER BY uploaded_at, document_id",
            (case_id,),
        )
        return [(row[0], row[1]) for row in cur.fetchall()]
