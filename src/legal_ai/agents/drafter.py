"""The drafting agent -- turn a conversation into a document.

The drafting prompt, and what settled each line of it.

One prompt for every document. An earlier version had one per instrument,
each with its own template and a rule deciding which conversations it
fitted; that answered "no document fits this thread" to almost everything,
because there is one prompt for every hundred documents Indian practice
uses. This one asks the model to choose the document the conversation calls
for and to lay it out in sections, which is what every legal document is.

The rules that remain are the ones that were earned. Each is here because
its absence produced a defect, measured against a real matter over three
rounds:

    a document that *is* a demand notice recited "a demand notice was sent
    on 2 August" among its facts, so it cited itself;

    a paragraph carried a citation and no substance -- the model had been
    given section titles rather than section text, and padded rather than
    omitted;

    the conversation's own findings were ignored in favour of the model's
    general knowledge;

    figures were restated from memory. `Kaveri Plastics v Mahdoom Bawa`
    (SC, 2025) holds a sum demanded in a s.138 notice must equal the cheque
    exactly, and a rupee's difference invalidates the notice -- so a figure
    is quoted as the conversation gave it or not written at all.

One model call, then deterministic checks. A second call could review the
draft, but every defect worth catching is mechanical -- a citation that was
never retrieved, a document reciting itself as already sent -- and a regex
catches those every time rather than most of the time. CLAUDE.md section 4.

It runs only when a reader asks for a document. It is not a stage of the
research graph: a document is not part of answering a question, and a node
that ran on every turn would draft one nobody asked for.
"""

from __future__ import annotations

import io
import json
import re
from datetime import date
from dataclasses import dataclass

from legal_ai.config import DEFAULT_CONFIG
from legal_ai.llm.client import generate
from legal_ai.schemas.draft import DraftStructure, Paragraph, Section

DRAFT_PROMPT = """You are an Indian legal drafter. Draft the document this
conversation calls for, from the conversation itself and the law it
established.

CHOOSE THE DOCUMENT. Nobody has told you what to draft. Read what was asked
and what was settled, and produce the document that follows from it -- a
legal notice, a written opinion, a reply, a bail application, an agreement,
a memo of advice. Name it in "title" as it would be headed on the page.
Where the conversation is a question of law with no matter behind it, an
opinion is what follows; where it is a client's grievance with a remedy in
sight, the instrument that pursues that remedy is.

LAY IT OUT IN SECTIONS. Every legal document is headings with numbered
paragraphs under them. Choose the headings this document needs -- FACTS,
THE POSITION, DEMAND, PRAYER, CONCLUSION, whatever it calls for -- and put
the paragraphs under them. Do not force a document into headings it does
not have.

Return ONLY this JSON:

{{"title": "as it would be headed on the page, in capitals",
  "subject": "one line saying what the document is about",
  "addressed_to": "the person it is addressed to, or empty",
  "on_behalf_of": "the client it is written for, or empty",
  "sections": [
    {{"heading": "FACTS",
      "paragraphs": [{{"text": "...", "authorities": []}}]}},
    {{"heading": "THE POSITION",
      "paragraphs": [{{"text": "...", "authorities": ["act:2189:sec-138"]}}]}}
  ],
  "needs_input": ["..."],
  "warnings": ["..."]}}

Rules:
- Facts are chronological and material only, with no legal argument in
  them. State only what the conversation actually gave. Where it asked a
  question in the abstract and supplied no facts, write no facts section
  rather than inventing a client's circumstances.
- One paragraph per proposition. A conversation that settled five things
  produces five paragraphs, each carrying the authority it rests on --
  merging them loses which authority settled which point.
- "authorities": document ids copied EXACTLY as they appear in RETRIEVED
  LAW, bare and without brackets, e.g. act:2189:sec-138. Cite nothing that
  is not in that list. Leave the array empty for a paragraph of fact.
- State the substance of a provision in your own words from the text you
  were given. If a provision does not bear on this document, omit it --
  never write a paragraph that carries a citation and says nothing.
- Where the document is itself the step being taken -- a notice being
  given, an application being made -- do not recite that step as already
  taken. The document would cite itself.
- Figures, dates and amounts are quoted exactly as the conversation gave
  them, or not written at all. Never restate a figure from memory and never
  round one. Where the document states a sum that must match an instrument
  exactly -- the amount of a cheque, of a decree, of an award -- add a
  "needs_input" item asking for it to be checked against the instrument
  itself: a sum that does not match can invalidate the document, and the
  conversation is not the instrument.
- Use what the conversation established -- its periods, deadlines and
  consequences -- rather than your own knowledge of the law.
- "needs_input": everything we do not hold that is needed before this can
  be used, including the advocate's own details.
- "warnings": anything making this the wrong document for this matter, or
  that must be resolved first -- a limitation period at risk, a fact in
  conflict with another. Empty array if there is nothing. Never invent one.
- Formal legal English, plain not archaic. Short sentences.
- A draft for an advocate to settle. Never invent an advocate's name,
  enrolment number, letterhead or seal, and never sign it.

MATTER
{matter}

CONVERSATION SO FAR
{conversation}

RETRIEVED LAW (the only authorities you may cite; text as held)
{law}
"""


# --- checks -----------------------------------------------------------
# A document reciting the step it is itself taking: a notice that says a
# notice was sent, an application that says an application was made. The
# document then cites itself.
_ALREADY_TAKEN = re.compile(
    r"\b(demand |statutory |legal )?(notice|application|complaint)\b[^.]{0,80}\b"
    r"(was|were|had been|has been)\s+(sent|issued|served|dispatched|filed|made)\b",
    re.IGNORECASE,
)


def validate(draft: DraftStructure, retrieved: set[str]) -> list[str]:
    """Every rule this draft breaks, in the order they are worth reading.

    `retrieved` is the set of document ids actually put in front of the
    drafter. A paragraph citing anything else is a fabricated citation,
    which is the worst failure this feature has.
    """
    failures: list[str] = []
    paragraphs = [p for block in draft.sections for p in block.paragraphs]

    for paragraph in paragraphs:
        for authority in paragraph.authorities:
            if authority not in retrieved:
                failures.append(
                    f"cites {authority!r}, which was not retrieved for this draft"
                )
        if paragraph.authorities and not paragraph.text.strip():
            failures.append("carries a citation with no text under it")

    for paragraph in paragraphs:
        if _ALREADY_TAKEN.search(paragraph.text):
            failures.append(
                "recites the step this document is itself taking as already "
                f"taken: {paragraph.text[:70]!r}"
            )
            break

    if not draft.title.strip():
        failures.append("has no title, so nothing says what it is")
    if not paragraphs:
        failures.append("has no paragraphs in it")
    elif not any(p.authorities for p in paragraphs):
        failures.append("rests on no law the conversation established")

    return failures

# --- the agent --------------------------------------------------------
@dataclass(frozen=True)
class DraftResult:
    structure: DraftStructure | None
    # Rules the draft breaks. Empty on a draft fit to render.
    failures: tuple[str, ...] = ()


def draft(matter: str, conversation: str, law: str, retrieved: set[str]) -> DraftResult:
    """A document drafted from one conversation, checked before it returns.

    `retrieved` is the set of ids actually put in front of the model. A
    paragraph citing anything else is a fabricated citation, which is the
    worst failure this feature has.
    """
    if not retrieved:
        return DraftResult(
            None, ("this conversation has not established any law to draft from",)
        )

    try:
        raw = generate(
            DRAFT_PROMPT.format(matter=matter, conversation=conversation, law=law),
            max_output_tokens=DEFAULT_CONFIG.draft_model_max_tokens,
        )
    except Exception:
        return DraftResult(None, ("the model was unreachable",))

    parsed = _parse(raw)
    if parsed is None:
        return DraftResult(None, ("the model's reply could not be read as a draft",))

    structure = _structure(parsed)
    return DraftResult(structure, tuple(validate(structure, retrieved)))


def _parse(raw: str) -> dict | None:
    cleaned = (raw or "").strip()
    if cleaned.startswith("```"):
        parts = cleaned.split("```")
        cleaned = parts[1] if len(parts) > 1 else cleaned
        if cleaned.startswith("json"):
            cleaned = cleaned[4:]
    try:
        parsed = json.loads(cleaned.strip())
    except ValueError:
        return None
    return parsed if isinstance(parsed, dict) else None


def _structure(parsed: dict) -> DraftStructure:
    return DraftStructure(
        title=_text(parsed.get("title")),
        subject=_text(parsed.get("subject")),
        addressed_to=_text(parsed.get("addressed_to")),
        on_behalf_of=_text(parsed.get("on_behalf_of")),
        sections=tuple(
            _section(block)
            for block in parsed.get("sections") or []
            if isinstance(block, dict)
        ),
        needs_input=tuple(_text(n) for n in parsed.get("needs_input") or [] if _text(n)),
        warnings=tuple(_text(w) for w in parsed.get("warnings") or [] if _text(w)),
    )


def _section(block: dict) -> Section:
    return Section(
        heading=_text(block.get("heading")),
        paragraphs=tuple(
            Paragraph(
                text=_text(item.get("text")),
                authorities=tuple(
                    _text(a) for a in item.get("authorities") or [] if _text(a)
                ),
            )
            for item in block.get("paragraphs") or []
            if isinstance(item, dict) and _text(item.get("text"))
        ),
    )


def _text(value) -> str:
    return str(value).strip() if value is not None else ""


# --- the file ---------------------------------------------------------
# Courts and opposing counsel read these on paper. 12pt serif with generous
# leading is the register; a UI font reads as a printout of a web page.
BODY_FONT = "Times New Roman"
BODY_SIZE_PT = 12

# What the advocate has to fill where we hold no value. Visible on purpose:
# a blank space reads as nothing missing.
UNFILLED = "__________"

FOOTER = (
    "DRAFT prepared by Pramāṇa AI from this matter's own record, on the "
    "authorities cited. To be settled and signed by an advocate. Verify "
    "every fact and provision before relying on it."
)


def render(
    draft: DraftStructure,
    *,
    today: date | None = None,
    citations: dict[str, str] | None = None,
) -> bytes:
    """The draft as a .docx, ready to download.

    `citations` maps an authority id to how it should be cited. Without one
    the body would print `act:2189:sec-138`, which is how this system
    addresses a provision and not how a document cites one -- prefer
    `render_with_citations`.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    citations = citations or {}
    document = Document()

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE_PT)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.4

    section = document.sections[0]
    section.left_margin = section.right_margin = Inches(1.1)
    section.top_margin = section.bottom_margin = Inches(1.0)

    def para(text="", *, bold=False, align=None):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = bold
        if align is not None:
            paragraph.alignment = align
        return paragraph

    para(draft.title or "DRAFT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    if draft.subject:
        para(draft.subject, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(
        f"Date: {(today or date.today()).strftime('%d %B %Y')}",
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )

    if draft.addressed_to:
        para("To,", bold=True)
        para(draft.addressed_to)
    if draft.on_behalf_of:
        para(f"On behalf of: {draft.on_behalf_of}")

    for block in draft.sections:
        if block.heading:
            para(block.heading.upper(), bold=True)
        for number, paragraph in enumerate(block.paragraphs, start=1):
            para(f"{number}. {_with_citations(paragraph, citations)}")

    para()
    para("_______________________________")
    para("Advocate")
    para("[Name, Bar Council enrolment number and address to be completed]")

    # Everything the draft could not settle, kept together at the end so an
    # advocate reads it once and deletes the block.
    if draft.warnings or draft.needs_input:
        document.add_page_break()
        para("DRAFTING NOTES — DELETE BEFORE SENDING", bold=True)
        if draft.warnings:
            para("Resolve before this is used:", bold=True)
            for warning in draft.warnings:
                para(f"• {warning}")
        if draft.needs_input:
            para("To be supplied:", bold=True)
            for item in draft.needs_input:
                para(f"• {item}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(FOOTER)
    run.font.size = Pt(8)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def render_with_citations(conn, draft: DraftStructure, *, today: date | None = None) -> bytes:
    """`render`, with every authority resolved against the corpus first."""
    from legal_ai.knowledge.static.citation import format_citation

    citations = {}
    for block in draft.sections:
        for paragraph in block.paragraphs:
            for authority in paragraph.authorities:
                if authority not in citations:
                    cited = format_citation(conn, authority)
                    if cited is not None:
                        citations[authority] = cited
    return render(draft, today=today, citations=citations)


def _with_citations(paragraph, citations: dict[str, str]) -> str:
    """The paragraph, with its authorities cited after it.

    An id the corpus cannot place is dropped rather than printed: a raw
    identifier reads as a bug to whoever receives the document, and the
    proposition still stands on the text around it.
    """
    cited = [citations[a] for a in paragraph.authorities if a in citations]
    if not cited:
        return paragraph.text
    return f"{paragraph.text} ({'; '.join(cited)})"
