"""The front door: an uploaded file becomes case facts.

Everything else in Phase 4 consumes DocumentFacts. Until this existed,
nothing produced them from a real file.
"""

import io

import pytest

from legal_ai.case.files import (
    UnsupportedDocument,
    extract_text,
    get_case_file_text,
    list_case_files,
)
from legal_ai.case.store import create_case, ensure_case_schema, get_case
from legal_ai.case.upload import document_id_for, upload_document
from legal_ai.knowledge.static.db import get_connection

COMPLAINT = """IN THE GUJARAT REAL ESTATE REGULATORY AUTHORITY
Kishor Patel, Complainant versus Marvel Developers Pvt Ltd, Respondent
The agreement to sell was executed on 12 March 2019.
Possession was due on 30 June 2021 and has not been handed over.
The complainant relies on Section 18 of the Real Estate (Regulation and
Development) Act, 2016.
"""


@pytest.fixture
def conn():
    connection = get_connection()
    ensure_case_schema(connection)
    with connection.cursor() as cur:
        cur.execute("DELETE FROM cases WHERE case_id LIKE 'test:%'")
    connection.commit()
    create_case(connection, "test:u1", "Patel v. Marvel")
    yield connection
    with connection.cursor() as cur:
        cur.execute("DELETE FROM cases WHERE case_id LIKE 'test:%'")
    connection.commit()
    connection.close()


def _docx_bytes(paragraphs, table_rows=()):
    import docx

    document = docx.Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=2)
        for row_index, (left, right) in enumerate(table_rows):
            table.cell(row_index, 0).text = left
            table.cell(row_index, 1).text = right
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- text extraction ---

def test_plain_text_is_read():
    assert "Kishor Patel" in extract_text(COMPLAINT.encode(), "complaint.txt")


def test_docx_paragraphs_are_read():
    data = _docx_bytes(["Possession was due on 30 June 2021."])
    assert "30 June 2021" in extract_text(data, "complaint.docx")


def test_docx_tables_are_read():
    # Legal agreements put substance in tables -- schedules, payment plans,
    # party details. Dropping them loses exactly the facts a case needs.
    data = _docx_bytes(["Preamble"], table_rows=[("Amount paid", "Rs 55,80,000")])
    text = extract_text(data, "complaint.docx")
    assert "Rs 55,80,000" in text


def test_an_unsupported_format_is_refused_by_name():
    with pytest.raises(UnsupportedDocument) as excinfo:
        extract_text(b"\x00\x01", "scan.tiff")
    assert ".tiff" in str(excinfo.value)


def test_a_file_with_no_extension_is_refused():
    with pytest.raises(UnsupportedDocument):
        extract_text(b"text", "noextension")


def test_extraction_is_capped():
    from legal_ai.case.files import MAX_TEXT_CHARS

    huge = ("x" * 1000 + "\n").encode() * 5000
    assert len(extract_text(huge, "big.txt")) <= MAX_TEXT_CHARS


# --- upload ---

def test_upload_stores_the_text_and_attaches_it_to_the_case(conn):
    facts = upload_document(conn, "test:u1", "complaint.txt", COMPLAINT.encode(),
                            extract_facts=False)
    assert get_case(conn, "test:u1").document_ids == (facts.document_id,)
    assert "Kishor Patel" in get_case_file_text(conn, facts.document_id)


def test_uploaded_files_do_not_enter_the_public_corpus(conn):
    # A client's pleading in `documents` would be returned by hybrid_search
    # as authority, and could surface for another user's query.
    facts = upload_document(conn, "test:u1", "complaint.txt", COMPLAINT.encode(),
                            extract_facts=False)
    with conn.cursor() as cur:
        cur.execute("SELECT count(*) FROM documents WHERE document_id = %s", (facts.document_id,))
        assert cur.fetchone()[0] == 0
        cur.execute("SELECT count(*) FROM document_chunks WHERE document_id = %s", (facts.document_id,))
        assert cur.fetchone()[0] == 0


def test_re_uploading_the_same_file_replaces_it(conn):
    # A user correcting a bad scan should not end up with two copies of one
    # exhibit that they then have to notice and delete.
    upload_document(conn, "test:u1", "complaint.txt", b"first version", extract_facts=False)
    upload_document(conn, "test:u1", "complaint.txt", b"corrected version", extract_facts=False)
    files = list_case_files(conn, "test:u1")
    assert len(files) == 1
    assert get_case_file_text(conn, files[0][0]) == "corrected version"


def test_the_same_filename_in_two_cases_is_two_documents(conn):
    create_case(conn, "test:u2", "Other matter")
    a = upload_document(conn, "test:u1", "notice.txt", b"one", extract_facts=False)
    b = upload_document(conn, "test:u2", "notice.txt", b"two", extract_facts=False)
    assert a.document_id != b.document_id


def test_an_unreadable_scan_is_stored_rather_than_rejected(conn):
    # A scanned PDF with no text layer is still an exhibit. The user needs
    # to see that we could not read it, not have the upload silently fail.
    facts = upload_document(conn, "test:u1", "scan.txt", b"   \n  ", extract_facts=True)
    assert facts.parties == () and facts.issues == ()
    assert get_case(conn, "test:u1").document_ids == (facts.document_id,)


def test_bulk_upload_can_skip_extraction(conn, monkeypatch):
    # Forty exhibits should list immediately, not wait on forty model calls.
    def boom(*a, **k):
        raise AssertionError("should not extract")

    monkeypatch.setattr("legal_ai.agents.document.extract_document_facts", boom)
    facts = upload_document(conn, "test:u1", "a.txt", COMPLAINT.encode(), extract_facts=False)
    assert facts.document_id


def test_extraction_runs_when_asked(conn, monkeypatch):
    from legal_ai.context.models import DocumentFacts

    monkeypatch.setattr(
        "legal_ai.case.upload.extract_document_facts",
        lambda document_id, text: DocumentFacts(document_id=document_id, issues=("delay",)),
    )
    facts = upload_document(conn, "test:u1", "a.txt", COMPLAINT.encode())
    assert facts.issues == ("delay",)


def test_the_document_id_is_stable_for_a_case_and_filename():
    assert document_id_for("c1", "notice.pdf") == document_id_for("c1", "notice.pdf")
    assert document_id_for("c1", "notice.pdf") != document_id_for("c2", "notice.pdf")


def test_deleting_a_case_deletes_its_files(conn):
    facts = upload_document(conn, "test:u1", "a.txt", COMPLAINT.encode(), extract_facts=False)
    conn.execute("DELETE FROM cases WHERE case_id = 'test:u1'")
    conn.commit()
    assert get_case_file_text(conn, facts.document_id) is None
