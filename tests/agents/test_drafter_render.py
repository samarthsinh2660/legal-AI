"""A drafted structure becomes a .docx someone can send.

Built directly rather than through a template file: a template per document
type is a promise to write one for every document Indian practice uses, and
the binaries git never tracked left the server with nothing to render with.
"""

import io
from datetime import date

from docx import Document

from legal_ai.schemas.draft import DraftStructure, Paragraph, Section
from legal_ai.agents.drafter import render


def _draft(**overrides) -> DraftStructure:
    base = dict(
        title="LEGAL OPINION",
        subject="Essential ingredients of criminal conspiracy",
        sections=(
            Section(
                "THE POSITION",
                (
                    Paragraph("An agreement between two or more persons is needed.",
                              ("act:ipc-1860:sec-120A",)),
                    Paragraph("The agreement may be express or implied."),
                ),
            ),
        ),
    )
    base.update(overrides)
    return DraftStructure(**base)


def _text(data: bytes) -> str:
    return "\n".join(p.text for p in Document(io.BytesIO(data)).paragraphs)


def test_the_document_carries_its_title_and_headings():
    body = _text(render(_draft()))

    assert "LEGAL OPINION" in body
    assert "THE POSITION" in body


def test_paragraphs_are_numbered_within_their_section():
    body = _text(render(_draft()))

    assert "1. An agreement between two or more persons is needed." in body
    assert "2. The agreement may be express or implied." in body


def test_a_raw_document_id_never_reaches_the_body():
    """`act:ipc-1860:sec-120A` is how this system addresses a provision,
    not how a document cites one."""
    body = _text(render(_draft(), citations={
        "act:ipc-1860:sec-120A": "Section 120A of the Indian Penal Code, 1860"
    }))

    assert "act:ipc-1860" not in body
    assert "Section 120A of the Indian Penal Code, 1860" in body


def test_an_unresolvable_citation_is_dropped_rather_than_printed():
    body = _text(render(_draft(), citations={}))

    assert "act:ipc-1860" not in body
    assert "An agreement between two or more persons is needed." in body


def test_warnings_and_missing_inputs_reach_the_document():
    body = _text(render(_draft(
        warnings=("The limitation period may have lapsed.",),
        needs_input=("Advocate's enrolment number",),
    )))

    assert "DELETE BEFORE SENDING" in body
    assert "may have lapsed" in body
    assert "enrolment number" in body


def test_a_clean_draft_carries_no_drafting_notes():
    assert "DELETE BEFORE SENDING" not in _text(render(_draft()))


def test_a_document_addressed_to_someone_says_so():
    body = _text(render(_draft(addressed_to="Mr. Rohan Malhotra")))

    assert "To," in body and "Mr. Rohan Malhotra" in body


def test_a_document_addressed_to_nobody_has_no_address_block():
    assert "To," not in _text(render(_draft()))


def test_the_date_is_the_day_it_was_drafted():
    assert "05 September 2026" in _text(render(_draft(), today=date(2026, 9, 5)))


def test_every_document_carries_the_draft_footer():
    from legal_ai.agents.drafter import FOOTER

    data = render(_draft())
    footer = Document(io.BytesIO(data)).sections[0].footer.paragraphs[0].text

    assert FOOTER[:40] in footer
